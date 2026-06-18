from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "safechat-partner-overview.md"
OUTPUT = ROOT / "docs" / "SafeChat-Partner-Overview.docx"
LOGO = ROOT / "docs" / "images" / "safechat-print-mark.png"

NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5B, 0x65, 0x73)
LIGHT = "F2F4F7"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, size=11, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), str(color).replace("RGBColor(", "").replace(")", "").replace(", ", ""))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([run_color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text):
    pattern = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|https?://[^\s]+|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), italic=True, color=MUTED)
        else:
            url = token.rstrip(".,);")
            suffix = token[len(url):]
            add_hyperlink(paragraph, url, url)
            if suffix:
                set_run_font(paragraph.add_run(suffix))
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]))


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Lead" not in [style.name for style in styles]:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(12)
    lead.font.color.rgb = NAVY
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.15


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, fmt, text, alignment, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(56)
        p.paragraph_format.space_after = Pt(30)
        picture = p.add_run().add_picture(str(LOGO), width=Inches(0.72))
        doc_pr = picture._inline.docPr
        doc_pr.set("descr", "SafeChat logo")
        doc_pr.set("title", "SafeChat")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(9)
    set_run_font(kicker.add_run("PARTNER & RESEARCH OVERVIEW"), size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("SafeChat"), size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run_font(
        subtitle.add_run("A privacy-preserving crisis-routing layer for AI conversations"),
        size=15,
        color=MUTED,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    set_run_font(
        p.add_run("  Sovereign local detection. Transparent routing. Human support.\nNo conversation text sent to SafeChat.  "),
        size=12,
        color=NAVY,
        bold=True,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("Interactive demonstration"), size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    add_hyperlink(p, "rob-e-graham.github.io/safechat/app/inspector.html", "https://rob-e-graham.github.io/safechat/app/inspector.html")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Rob Graham  |  FAMTEC  |  18 June 2026"), size=10, color=MUTED)
    doc.add_page_break()


def add_body_from_markdown(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    skip_headings = 2
    list_kind = None
    list_num_id = None
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(line.strip() for line in paragraph_buffer)
        style = "Lead" if text.startswith("SafeChat is a source-available") else None
        p = doc.add_paragraph(style=style)
        add_inline(p, text)
        paragraph_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") or stripped.startswith("## A privacy-preserving"):
            if skip_headings:
                skip_headings -= 1
                continue
        if not stripped:
            flush_paragraph()
            list_kind = None
            list_num_id = None
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(p, heading.group(2))
            list_kind = None
            list_num_id = None
            continue
        item = re.match(r"^(-|\d+\.)\s+(.+)$", stripped)
        if item:
            flush_paragraph()
            kind = "bullet" if item.group(1) == "-" else "decimal"
            if list_kind != kind or list_num_id is None:
                list_kind = kind
                list_num_id = add_numbering(doc, kind)
            p = doc.add_paragraph()
            apply_numbering(p, list_num_id)
            add_inline(p, item.group(2))
            continue
        paragraph_buffer.append(stripped.rstrip("  "))
    flush_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    doc.core_properties.title = "SafeChat Partner and Research Overview"
    doc.core_properties.subject = "Purpose, operation, VERA-MH integration, research, licensing, and development opportunities"
    doc.core_properties.author = "Rob Graham / FAMTEC"
    doc.core_properties.keywords = "SafeChat, VERA-MH, AI safety, crisis routing, research"

    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("SAFECHAT  |  PARTNER & RESEARCH OVERVIEW"), size=9, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)

    add_cover(doc)
    add_body_from_markdown(doc)

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            run.font.name = run.font.name or "Calibri"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
